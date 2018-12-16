import os
import sys
import io
import docx
import random
import numpy as np

input_file = r'D:\student\英语\Proficiency Test 2018 - 副本.docx'
output_file = r'D:\student\英语\Proficiency Test 2018_result.docx'
data = []
doc = docx.Document(input_file)
for item in doc.paragraphs:
    data.append(item.text)


file = docx.Document()
for i in range(len(data)):
    file.add_paragraph(data[i])
file.save(output_file)
print('success')